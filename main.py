import flet as ft
import asyncio

from datetime import datetime
import sqlite3
import os

from reportlab.pdfgen import canvas
from docx import Document
from docx.shared import Pt

VERMELHO = "#C00000"
VERMELHO_ESCURO = "#8B0000"
VERMELHO_HOVER = "#A30000"
BRANCO = "#FFFFFF"
CINZA_FUNDO = "#F3F4F6"
CINZA_CARD = "#FFFFFF"
CINZA_BORDA = "#E2E4E8"
TEXTO_PRIMARIO = "#1F2024"
TEXTO_SECUNDARIO = "#6B7280"
SIDEBAR_BG = "#7A0000"
SIDEBAR_ITEM_HOVER = "#931414"
SIDEBAR_ITEM_ATIVO = "#A30000"

FONTE_BASE = "Segoe UI"


def caminho_banco():
    base = os.environ.get("FLET_APP_STORAGE_DATA")
    if base:
        return os.path.join(base, "orcamentos.db")
    return os.path.join(os.path.dirname(os.path.abspath(__file__)), "orcamentos.db")


DB_PATH = caminho_banco()


CARACTERISTICAS_ELEVADOR_SOCIAL = (
    "Tipo: Sem Casa de Maquinas. Maquina: Sincrona Gearless. "
    "Capacidade para 08 passageiros - 630Kg. Velocidade: 1,5m/s. "
    "Dimensoes internas da cabine: 1.100L x 1.400P x 2.200mm."
)

PARCELAS = [
    "a vista", "parcela 2x", "parcela 3x", "parcela 4x", "parcela 5x",
    "parcela 6x", "parcela 7x", "parcela 8x", "parcela 9x",
    "parcela 10x", "parcela 11x", "parcela 12x",
]


def formatar_moeda(valor):
    valor = float(valor)
    return f"R$ {valor:,.2f}".replace(",", "X").replace(".", ",").replace("X", ".")


def formatar_moeda_seguro(valor):
    """Como formatar_moeda, mas nunca quebra a tela se o valor salvo no banco
    estiver corrompido (ex.: registros antigos com dado inconsistente)."""
    try:
        return formatar_moeda(valor)
    except (ValueError, TypeError):
        return "valor inválido"


def para_float_seguro(valor):
    """Converte para float com fallback 0.0 caso o dado esteja corrompido."""
    try:
        return float(valor)
    except (ValueError, TypeError):
        return 0.0


def caracteristicas_produto(nome, caracteristicas):
    """Mantem as caracteristicas padrao do Elevador Social se o banco antigo estiver vazio."""
    texto = (caracteristicas or "").strip()
    if nome == "Elevador Social" and not texto:
        return CARACTERISTICAS_ELEVADOR_SOCIAL
    return texto


def _nome_normalizado(nome):
    return (nome or "").strip().lower()


def _eh_elevador_social(nome):
    return _nome_normalizado(nome) == "elevador social"


def _eh_elevador_maca(nome):
    return "maca" in _nome_normalizado(nome)


def _dois_digitos(valor):
    return f"{int(valor):02d}"


def _resumo_quantidade_word(itens):
    qtd_social = 0
    qtd_maca = 0
    outros = {}
    ordem_outros = []
    for item in itens:
        produto = item["produto"]
        quantidade = item["quantidade"]
        if _eh_elevador_social(produto):
            qtd_social += quantidade
        elif _eh_elevador_maca(produto):
            qtd_maca += quantidade
        else:
            if produto not in outros:
                outros[produto] = 0
                ordem_outros.append(produto)
            outros[produto] += quantidade

    partes = []
    if qtd_social:
        nome = "Elevador Social" if qtd_social == 1 else "Elevadores Sociais"
        partes.append(f"{_dois_digitos(qtd_social)} x {nome}")
    if qtd_maca:
        nome = "Elevador Maca" if qtd_maca == 1 else "Elevadores Maca"
        partes.append(f"{_dois_digitos(qtd_maca)} {nome}")
    for produto in ordem_outros:
        partes.append(f"{_dois_digitos(outros[produto])} x {produto}")
    return "Quantidade: " + " e ".join(partes)




def _adicionar_paragrafo(doc, texto, negrito=False, tamanho=12):
    paragrafo = doc.add_paragraph()
    run = paragrafo.add_run(texto)
    run.bold = negrito
    run.font.size = Pt(tamanho)
    return paragrafo


def gerar_proposta_word_liftdesk(caminho_saida, dados):
    """Gera a proposta em Word do zero, sem depender de template externo."""
    doc = Document()

    _adicionar_paragrafo(doc, "Proposta Comercial - LiftDesk", negrito=True, tamanho=18)
    _adicionar_paragrafo(doc, f"Cliente: {dados['cliente']}")
    _adicionar_paragrafo(doc, f"Telefone: {dados['telefone']}")
    _adicionar_paragrafo(doc, f"E-mail: {dados['email']}")
    _adicionar_paragrafo(doc, f"Data/Hora: {dados['data_hora']}")
    doc.add_paragraph()

    _adicionar_paragrafo(doc, "Itens do orçamento", negrito=True)
    for item in dados["itens"]:
        _adicionar_paragrafo(
            doc,
            f"- {item['quantidade']}x {item['produto']} | "
            f"{formatar_moeda(item['preco_unitario'])} cada",
        )
    doc.add_paragraph()

    _adicionar_paragrafo(doc, f"Resumo: {_resumo_quantidade_word(dados['itens'])}")
    _adicionar_paragrafo(doc, f"Condição de pagamento: {dados['parcela']}")
    _adicionar_paragrafo(doc, f"Valor total: {dados['valor_total_formatado']}")
    _adicionar_paragrafo(doc, f"Valor da parcela: {dados['valor_parcela_formatado']}")

    doc.save(caminho_saida)
    return doc


def gerar_contrato_word_liftdesk(caminho_saida, placeholders):
    """Gera o contrato em Word do zero com os dados informados."""
    doc = Document()

    _adicionar_paragrafo(doc, "Contrato de Venda - LiftDesk", negrito=True, tamanho=18)
    _adicionar_paragrafo(doc, f"Cliente: {placeholders.get('cliente', '')}")
    _adicionar_paragrafo(doc, f"Telefone: {placeholders.get('telefone', '')}")
    _adicionar_paragrafo(doc, f"E-mail: {placeholders.get('email', '')}")
    doc.add_paragraph()

    for chave, valor in placeholders.items():
        _adicionar_paragrafo(doc, f"{chave}: {valor}")

    doc.save(caminho_saida)


def _limpar_e_escrever(paragrafo, novo_texto):
    """Limpa o conteúdo de um parágrafo e escreve novo texto."""
    paragrafo.clear()
    paragrafo.add_run(novo_texto)


def _localizar_paragrafo(doc, texto_procurado):
    """Localiza o primeiro parágrafo que contém o texto procurado."""
    for p in doc.paragraphs:
        if texto_procurado in p.text:
            return p
    return None


def _remover_paragrafo(paragrafo):
    """Remove um parágrafo do documento."""
    if paragrafo is not None:
        p = paragrafo._element
        p.getparent().remove(p)


def _inserir_paragrafo_apos(paragrafo_ref, novo_texto, copiar_estilo_de=None):
    """Insere um novo parágrafo após o parágrafo de referência."""
    if paragrafo_ref is None:
        return None
    novo_paragrafo = paragrafo_ref.insert_paragraph_before(novo_texto)
    if copiar_estilo_de:
        novo_paragrafo.style = copiar_estilo_de.style
    return novo_paragrafo


def gerar_proposta_word(caminho_saida, dados):
    """Gera a proposta em Word do zero, sem depender de template externo."""
    gerar_proposta_word_liftdesk(caminho_saida, dados)



def gerar_contrato_word(caminho_saida, placeholders):
    """Gera o contrato em Word do zero com os dados informados."""
    gerar_contrato_word_liftdesk(caminho_saida, placeholders)


def main(page: ft.Page):
    page.title = "LiftDesk — Sistema de Orçamentos"
    page.bgcolor = CINZA_FUNDO
    page.window.width = 1200
    page.window.height = 720
    page.window.min_width = 1000
    page.window.min_height = 640
    page.theme_mode = ft.ThemeMode.LIGHT
    page.theme = ft.Theme(font_family=FONTE_BASE)
    page.padding = 0

    # --------------------------------------------------------
    # BANCO DE DADOS
    # --------------------------------------------------------
    os.makedirs(os.path.dirname(DB_PATH), exist_ok=True)
    conexao = sqlite3.connect(DB_PATH)
    cursor = conexao.cursor()
    cursor.execute(
        """
        CREATE TABLE IF NOT EXISTS orcamentos (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            data_hora TEXT,
            cliente TEXT,
            email TEXT,
            telefone TEXT,
            parcela TEXT,
            valor_total REAL,
            valor_parcela REAL
        )
        """
    )
    cursor.execute(
        """
        CREATE TABLE IF NOT EXISTS orcamento_itens (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            orcamento_id INTEGER NOT NULL,
            produto TEXT,
            quantidade INTEGER,
            preco_unitario REAL,
            caracteristicas TEXT,
            FOREIGN KEY (orcamento_id) REFERENCES orcamentos (id)
        )
        """
    )
    cursor.execute(
        """
        CREATE TABLE IF NOT EXISTS produtos (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            nome TEXT UNIQUE,
            preco REAL,
            caracteristicas TEXT
        )
        """
    )
    cursor.execute(
        """
        CREATE TABLE IF NOT EXISTS contratos (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            orcamento_id INTEGER NOT NULL,
            cliente TEXT,
            cnpj TEXT,
            data_hora TEXT,
            FOREIGN KEY (orcamento_id) REFERENCES orcamentos (id)
        )
        """
    )

    # Migração defensiva: se a tabela orcamentos ou produtos já existir
    # de uma versão anterior (sem essas colunas), adiciona o que falta.
    cursor.execute("PRAGMA table_info(orcamentos)")
    colunas_orcamentos = {linha[1] for linha in cursor.fetchall()}
    if "telefone" not in colunas_orcamentos:
        cursor.execute("ALTER TABLE orcamentos ADD COLUMN telefone TEXT")
    if "valor_parcela" not in colunas_orcamentos:
        cursor.execute("ALTER TABLE orcamentos ADD COLUMN valor_parcela REAL")
    
    cursor.execute("PRAGMA table_info(produtos)")
    colunas_produtos = {linha[1] for linha in cursor.fetchall()}
    if "caracteristicas" not in colunas_produtos:
        cursor.execute("ALTER TABLE produtos ADD COLUMN caracteristicas TEXT")

    conexao.commit()

    # Popula a tabela de produtos com os itens originais, apenas
    # se ainda estiver vazia (primeira execução após a atualização).
    cursor.execute("SELECT COUNT(*) FROM produtos")
    if cursor.fetchone()[0] == 0:
        produtos_iniciais = [
            ("grua torre", 250000, ""),
            ("andaime suspenso", 180000, ""),
            ("Elevador Social", 180000, CARACTERISTICAS_ELEVADOR_SOCIAL),
            ("elevador panoramico", 120000, ""),
        ]
        cursor.executemany(
            "INSERT INTO produtos (nome, preco, caracteristicas) VALUES (?, ?, ?)",
            produtos_iniciais,
        )
        conexao.commit()

    cursor.execute(
        """
        UPDATE produtos
        SET caracteristicas = ?
        WHERE nome = ?
          AND (caracteristicas IS NULL OR TRIM(caracteristicas) = '')
        """,
        (CARACTERISTICAS_ELEVADOR_SOCIAL, "Elevador Social"),
    )
    conexao.commit()

    def carregar_produtos():
        """Lê a tabela produtos e devolve um dict {nome: {preco, caracteristicas}}."""
        cursor.execute("SELECT nome, preco, caracteristicas FROM produtos ORDER BY nome")
        return {
            nome: {"preco": preco, "caracteristicas": caracteristicas_produto(nome, caracteristicas)}
            for nome, preco, caracteristicas in cursor.fetchall()
        }

    def produto_existe(nome):
        cursor.execute("SELECT 1 FROM produtos WHERE nome = ?", (nome,))
        return cursor.fetchone() is not None

    def inserir_produto(nome, preco, caracteristicas=""):
        cursor.execute(
            "INSERT INTO produtos (nome, preco, caracteristicas) VALUES (?, ?, ?)",
            (nome, preco, caracteristicas),
        )
        conexao.commit()

    def excluir_produto(nome):
        cursor.execute("DELETE FROM produtos WHERE nome = ?", (nome,))
        conexao.commit()

    def carregar_lista_orcamentos():
        """Retorna todos os orçamentos salvos, do mais recente para o mais antigo,
        para popular o seletor da tela de Contratos."""
        cursor.execute(
            "SELECT id, data_hora, cliente, valor_total FROM orcamentos ORDER BY id DESC"
        )
        return cursor.fetchall()

    def carregar_orcamento_completo(orcamento_id):
        """Recarrega um orçamento salvo (cabeçalho + itens) no mesmo formato
        de 'dados' usado para gerar a proposta/contrato em Word."""
        cursor.execute(
            """SELECT id, data_hora, cliente, email, telefone, parcela,
                      valor_total, valor_parcela
               FROM orcamentos WHERE id = ?""",
            (orcamento_id,),
        )
        linha = cursor.fetchone()
        if not linha:
            return None
        (id_, data_hora, cliente, email, telefone, parcela,
         valor_total, valor_parcela) = linha

        cursor.execute(
            """SELECT produto, quantidade, preco_unitario, caracteristicas
               FROM orcamento_itens WHERE orcamento_id = ?""",
            (orcamento_id,),
        )
        itens = []
        for produto, quantidade, preco_unitario, caracteristicas in cursor.fetchall():
            preco_unitario = para_float_seguro(preco_unitario)
            quantidade = int(quantidade) if str(quantidade).strip().isdigit() else 0
            itens.append(
                {
                    "produto": produto,
                    "quantidade": quantidade,
                    "preco_unitario": preco_unitario,
                    "subtotal": preco_unitario * quantidade,
                    "caracteristicas": caracteristicas or "",
                }
            )

        valor_total = para_float_seguro(valor_total)
        valor_parcela = para_float_seguro(valor_parcela)

        return {
            "id": id_,
            "data_hora": data_hora,
            "cliente": cliente,
            "email": email,
            "telefone": telefone,
            "parcela": parcela,
            "valor_total": valor_total,
            "valor_parcela": valor_parcela,
            "valor_total_formatado": formatar_moeda_seguro(valor_total),
            "valor_parcela_formatado": formatar_moeda_seguro(valor_parcela),
            "itens": itens,
        }

    def registrar_contrato(orcamento_id, cliente, cnpj):
        cursor.execute(
            "INSERT INTO contratos (orcamento_id, cliente, cnpj, data_hora) VALUES (?,?,?,?)",
            (orcamento_id, cliente, cnpj, datetime.now().strftime("%d/%m/%Y, %H:%M:%S")),
        )
        conexao.commit()

    historico = []

    file_picker = ft.FilePicker()
    page.services.append(file_picker)

    def mostrar_snack(mensagem, erro=False):
        page.show_dialog(
            ft.SnackBar(
                content=ft.Text(mensagem, color=BRANCO),
                bgcolor=VERMELHO_ESCURO if erro else "#2E7D32",
            )
        )

    cliente_field = ft.TextField(
        label="Cliente",
        hint_text="Nome completo do cliente",
        height=56,
        border_radius=8,
        border_color=CINZA_BORDA,
        bgcolor="#FAFAFA",
        color=TEXTO_PRIMARIO,
        label_style=ft.TextStyle(color=TEXTO_SECUNDARIO),
        hint_style=ft.TextStyle(color=TEXTO_SECUNDARIO),
    )

    email_field = ft.TextField(
        label="E-mail",
        hint_text="cliente@email.com",
        height=56,
        border_radius=8,
        border_color=CINZA_BORDA,
        bgcolor="#FAFAFA",
        color=TEXTO_PRIMARIO,
        label_style=ft.TextStyle(color=TEXTO_SECUNDARIO),
        hint_style=ft.TextStyle(color=TEXTO_SECUNDARIO),
    )

    telefone_field = ft.TextField(
        label="Telefone",
        hint_text="(81) 99999-9999",
        height=56,
        border_radius=8,
        border_color=CINZA_BORDA,
        bgcolor="#FAFAFA",
        color=TEXTO_PRIMARIO,
        label_style=ft.TextStyle(color=TEXTO_SECUNDARIO),
        hint_style=ft.TextStyle(color=TEXTO_SECUNDARIO),
    )
    #obra_field = ft.TextField(
        #label= "Obra",
       # hint_text= "Nome da Obra",
        #height=56,
       # border_radius=8,
       # border_color=CINZA_BORDA,
       # bgcolor="#FAFAFA",
       # color=TEXTO_PRIMARIO,
       # label_style=ft.TextStyle(color=TEXTO_SECUNDARIO),
       # hint_style=ft.TextStyle(color=TEXTO_SECUNDARIO)
   # )
    # --------------------------------------------------------
    # LINHAS DINÂMICAS DE PRODUTO (lista de itens do orçamento)
    # ------------------------------------------------------
    # Cada linha é um dict: {"id": int, "dropdown": Dropdown,
    # "quantidade_field": TextField, "container": Container}
    # guardado em itens_produto_linhas, na ordem em que aparecem.
    # --------------------------------------------------------
    itens_produto_linhas = []
    proximo_id_linha = {"valor": 0}
    itens_produto_coluna = ft.Column(spacing=10)

    def criar_dropdown_produto():
        dd = ft.Dropdown(
            label="Produto",
            options=[],
            border_radius=8,
            bgcolor=VERMELHO,
            color=TEXTO_PRIMARIO,
            height=56,
        )
        produtos_atuais = carregar_produtos()
        dd.options = [
            ft.DropdownOption(key=nome, text=nome) for nome in produtos_atuais.keys()
        ]
        if produtos_atuais:
            dd.value = next(iter(produtos_atuais))
        return dd

    def remover_linha_produto(id_linha):
        def handler(e):
            nonlocal itens_produto_linhas
            itens_produto_linhas = [
                linha for linha in itens_produto_linhas if linha["id"] != id_linha
            ]
            redesenhar_linhas_produto()
            page.update()
        return handler

    def adicionar_linha_produto(e=None):
        id_linha = proximo_id_linha["valor"]
        proximo_id_linha["valor"] += 1

        dropdown = criar_dropdown_produto()
        qtd_field = ft.TextField(
            label="Quantidade",
            hint_text="0",
            height=56,
            border_radius=8,
            border_color=CINZA_BORDA,
            bgcolor="#FAFAFA",
            keyboard_type=ft.KeyboardType.NUMBER,
            color=TEXTO_PRIMARIO,
            label_style=ft.TextStyle(color=TEXTO_SECUNDARIO),
            hint_style=ft.TextStyle(color=TEXTO_SECUNDARIO),
        )

        linha = {"id": id_linha, "dropdown": dropdown, "quantidade_field": qtd_field}
        itens_produto_linhas.append(linha)
        redesenhar_linhas_produto()
        if e is not None:
            page.update()

    def redesenhar_linhas_produto():
        """Reconstrói a coluna visual a partir de itens_produto_linhas."""
        linhas_visuais = []
        for linha in itens_produto_linhas:
            pode_remover = len(itens_produto_linhas) > 1
            linhas_visuais.append(
                ft.Row(
                    controls=[
                        ft.Container(content=linha["dropdown"], expand=2),
                        ft.Container(content=linha["quantidade_field"], expand=1),
                        ft.IconButton(
                            icon=ft.Icons.DELETE_OUTLINE,
                            icon_color=VERMELHO if pode_remover else CINZA_BORDA,
                            tooltip="Remover produto" if pode_remover else None,
                            disabled=not pode_remover,
                            on_click=remover_linha_produto(linha["id"]),
                        ),
                    ],
                    spacing=10,
                )
            )
        itens_produto_coluna.controls = linhas_visuais

    def repopular_linhas_produto():
        """Recarrega as opções de produto em todas as linhas existentes
        (chamado ao abrir a tela de orçamento, caso produtos tenham mudado)."""
        if not itens_produto_linhas:
            adicionar_linha_produto()
            return
        for linha in itens_produto_linhas:
            produtos_atuais = carregar_produtos()
            linha["dropdown"].options = [
                ft.DropdownOption(key=nome, text=nome) for nome in produtos_atuais.keys()
            ]
            if produtos_atuais and linha["dropdown"].value not in produtos_atuais:
                linha["dropdown"].value = next(iter(produtos_atuais))
        redesenhar_linhas_produto()

    parcela_dropdown = ft.Dropdown(
        label="Condição de Pagamento",
        value="a vista",
        options=[ft.DropdownOption(key=p, text=p) for p in PARCELAS],
        border_radius=8,
        bgcolor=VERMELHO,
        color=TEXTO_PRIMARIO,
        height=56,
    )

    erro_text = ft.Container(
        content=ft.Text("", color=VERMELHO, weight=ft.FontWeight.BOLD, size=12),
        bgcolor="#FCEAEA",
        border_radius=8,
        padding=8,
        visible=False,
    )

    resultado_text = ft.Text(
        "Valor Total: R$ 0",
        size=14,
        color=TEXTO_PRIMARIO,
        selectable=True,
    )

    def mostrar_erro(mensagem):
        erro_text.content.value = mensagem
        erro_text.visible = True
        page.update()

    def esconder_erro():
        erro_text.visible = False

    def calcular(e=None):
        cliente = cliente_field.value or ""
        #obra = obra_field.value or ""
        produtos_atuais = carregar_produtos()

        if not itens_produto_linhas:
            mostrar_erro("Adicione ao menos um produto")
            return None

        itens_calculados = []
        for linha in itens_produto_linhas:
            produto = linha["dropdown"].value
            if not produto:
                mostrar_erro("Selecione um produto em todas as linhas")
                return None
            if produto not in produtos_atuais:
                mostrar_erro(f'Produto "{produto}" não existe mais — escolha outro')
                return None

            texto_quantidade = linha["quantidade_field"].value or ""
            try:
                quantidade = int(texto_quantidade)
                if quantidade <= 0:
                    raise ValueError
            except ValueError:
                mostrar_erro("Digite uma quantidade válida em todas as linhas!")
                return None

            preco_unitario = produtos_atuais[produto]["preco"]
            caracteristicas = caracteristicas_produto(
                produto, produtos_atuais[produto]["caracteristicas"]
            )
            itens_calculados.append(
                {
                    "produto": produto,
                    "quantidade": quantidade,
                    "preco_unitario": preco_unitario,
                    "subtotal": preco_unitario * quantidade,
                    "caracteristicas": caracteristicas,
                }
            )

        esconder_erro()

        if cliente.strip() == "":
            mostrar_erro("Digite o nome do cliente")
            return None

        email = email_field.value or ""
        if "@" not in email or "." not in email:
            mostrar_erro("E-mail inválido!")
            return None

        telefone = telefone_field.value or ""
       
        valor_total = sum(item["subtotal"] for item in itens_calculados)
        parcela = parcela_dropdown.value

        if parcela == "a vista":
            valor_parcela = valor_total
        else:
            numero_parcelas = int(parcela.split()[1].replace("x", ""))
            valor_parcela = valor_total / numero_parcelas

        valor_total_formatado = formatar_moeda(valor_total)
        valor_parcela_formatado = formatar_moeda(valor_parcela)

        resumo_produtos = "\n".join(
            f"{item['quantidade']}x {item['produto']}" for item in itens_calculados
        )
        resultado_text.value = (
            f"Cliente: {cliente}\n"
           #add obra!!
            f"{resumo_produtos}\n"
            f"Valor Total: {valor_total_formatado}\n"
            f"Valor Parcela: {valor_parcela_formatado}\n"
        )

        data_hora = datetime.now().strftime("%d/%m/%Y, %H:%M:%S")
        dados = {
            "cliente": cliente,
            #add obra!!
            "email": email,
            "telefone": telefone,
            "itens": itens_calculados,
            "parcela": parcela,
            "valor_total": valor_total,
            "valor_parcela": valor_parcela,
            "valor_parcela_formatado": valor_parcela_formatado,
            "valor_total_formatado": valor_total_formatado,
            "data_hora": data_hora,
        }

        historico.append(dados)

        cursor.execute(
            """INSERT INTO orcamentos (
                data_hora, cliente, email, telefone, parcela, valor_total, valor_parcela
            ) VALUES (?,?,?,?,?,?,?)""",
            (
                dados["data_hora"],
                dados["cliente"],
                dados["email"],
                #add obra!
                dados["telefone"],
                dados["parcela"],
                dados["valor_total"],
                dados["valor_parcela"],
            ),
        )
        orcamento_id = cursor.lastrowid
        for item in itens_calculados:
            cursor.execute(
                """INSERT INTO orcamento_itens (
                    orcamento_id, produto, quantidade, preco_unitario, caracteristicas
                ) VALUES (?,?,?,?,?)""",
                (
                    orcamento_id,
                    item["produto"],
                    item["quantidade"],
                    item["preco_unitario"],
                    item["caracteristicas"],
                ),
            )
        conexao.commit()

        dados["id"] = orcamento_id
        page.update()
        return dados

    def _escrever_txt(caminho, dados):
        linhas_produtos = "\n".join(
            f"  {item['quantidade']}x {item['produto']} — "
            f"{formatar_moeda(item['preco_unitario'])} cada"
            for item in dados["itens"]
        )
        conteudo = (
            f"=== ORÇAMENTO LIFTDESK ===\n\n"
            f"Data/Hora: {dados['data_hora']}\n\n"
            #add obra!!
            f"Cliente: {dados['cliente']}\n"
            f"E-mail: {dados['email']}\n"
            f"Telefone: {dados['telefone']}\n\n"
            f"Produtos:\n{linhas_produtos}\n\n"
            f"Condição de Pagamento: {dados['parcela']}\n"
            f"--------------------------------\n"
            f"VALOR TOTAL: {dados['valor_total_formatado']}\n"
            f"VALOR DA PARCELA: {dados['valor_parcela_formatado']}\n"
        )
        with open(caminho, "w", encoding="utf-8") as arquivo:
            arquivo.write(conteudo)

    def _escrever_pdf(caminho, dados):
        pdf = canvas.Canvas(caminho)
        pdf.setTitle("Orçamento LiftDesk")
        pdf.drawString(100, 800, "LiftDesk")
        pdf.drawString(100, 780, f"Data: {dados['data_hora']}")
       # pdf.drawString(100, 760, f"Obra: {dados['obra']}")
        pdf.drawString(100, 740, f"Cliente: {dados['cliente']}")
        pdf.drawString(100, 720, f"E-mail: {dados['email']}")
        pdf.drawString(100, 700, f"Telefone: {dados['telefone']}")

        y = 670
        pdf.drawString(100, y, "Produtos:")
        for item in dados["itens"]:
            y -= 18
            pdf.drawString(
                120,
                y,
                f"{item['quantidade']}x {item['produto']} — "
                f"{formatar_moeda(item['preco_unitario'])} cada",
            )

        y -= 30
        pdf.drawString(100, y, f"Pagamento: {dados['parcela']}")
        y -= 20
        pdf.drawString(100, y, f"Valor Total: {dados['valor_total_formatado']}")
        y -= 18
        pdf.drawString(100, y, f"Valor Parcela: {dados['valor_parcela_formatado']}")
        y -= 60
        pdf.drawString(100, y, "_________________")
        y -= 18
        pdf.drawString(100, y, "Assinatura")
        pdf.save()

    def _escrever_docx(caminho, dados):
        gerar_proposta_word(caminho, dados)

    async def salvar_orcamento_txt(e):
        dados = calcular()
        if not dados:
            return
        caminho = await file_picker.save_file(
            dialog_title="Salvar orçamento",
            file_name="orcamento.txt",
            allowed_extensions=["txt"],
        )
        if not caminho:
            return  # usuário cancelou
        if not caminho.lower().endswith(".txt"):
            caminho += ".txt"
        _escrever_txt(caminho, dados)
        mostrar_snack(f"Arquivo salvo em: {caminho}")

    async def salvar_pdf(e):
        dados = calcular()
        if not dados:
            return
        caminho = await file_picker.save_file(
            dialog_title="Salvar orçamento em PDF",
            file_name="orcamento.pdf",
            allowed_extensions=["pdf"],
        )
        if not caminho:
            return
        if not caminho.lower().endswith(".pdf"):
            caminho += ".pdf"
        _escrever_pdf(caminho, dados)
        mostrar_snack(f"Arquivo salvo em: {caminho}")

    async def salvar_word(e):
        dados = calcular()
        if not dados:
            return
        caminho = await file_picker.save_file(
            dialog_title="Salvar orçamento em Word",
            file_name="orcamento.docx",
            allowed_extensions=["docx"],
        )
        if not caminho:
            return
        if not caminho.lower().endswith(".docx"):
            caminho += ".docx"
        try:
            _escrever_docx(caminho, dados)
        except (FileNotFoundError, ValueError) as erro:
            mostrar_snack(str(erro), erro=True)
            return
        mostrar_snack(f"Arquivo salvo em: {caminho}")

    def botao_acao(texto, on_click, principal=False):
        return ft.ElevatedButton(
            content=texto,
            on_click=on_click,
            height=42,
            style=ft.ButtonStyle(
                bgcolor=VERMELHO if principal else BRANCO,
                color=BRANCO if principal else TEXTO_PRIMARIO,
                shape=ft.RoundedRectangleBorder(radius=10),
                side=ft.BorderSide(0 if principal else 1, CINZA_BORDA),
            ),
        )

    barra_acoes = ft.Row(
        controls=[
            botao_acao("Calcular Orçamento", calcular, principal=True),
            botao_acao("Salvar TXT", salvar_orcamento_txt),
            botao_acao("Salvar PDF", salvar_pdf),
            botao_acao("Salvar Word", salvar_word),
            botao_acao("Histórico", lambda e: asyncio.create_task(page.push_route("/historico"))),
        ],
        wrap=True,
        spacing=10,
    )

    botao_add_produto = ft.TextButton(
        content=ft.Row(
            controls=[
                ft.Icon(ft.Icons.ADD, color=VERMELHO, size=18),
                ft.Text("Adicionar Produto", color=VERMELHO, weight=ft.FontWeight.BOLD, size=13),
            ],
            spacing=6,
            tight=True,
        ),
        on_click=adicionar_linha_produto,
    )

    card_formulario = ft.Container(
        content=ft.Column(
            controls=[
                cliente_field,
                email_field,
                telefone_field,
                ft.Text(
                    "PRODUTOS",
                    size=11,
                    weight=ft.FontWeight.BOLD,
                    color=TEXTO_SECUNDARIO,
                ),
                itens_produto_coluna,
                ft.Row(controls=[botao_add_produto], alignment=ft.MainAxisAlignment.START),
                parcela_dropdown,
                erro_text,
            ],
            spacing=16,
        ),
        bgcolor=CINZA_CARD,
        border_radius=14,
        border=ft.Border.all(1, CINZA_BORDA),
        padding=26,
    )

    card_resultado = ft.Container(
        content=ft.Column(
            controls=[
                ft.Text(
                    "RESUMO DO ORÇAMENTO",
                    size=12,
                    weight=ft.FontWeight.BOLD,
                    color=TEXTO_SECUNDARIO,
                ),
                resultado_text,
            ],
            spacing=14,
            
        ),
        bgcolor=CINZA_CARD,
        border_radius=14,
        border=ft.Border.all(1, CINZA_BORDA),
        padding=24,
    )
   
        
    # --------------------------------------------------------
    # VIEW PRINCIPAL — NOVO ORÇAMENTO
    # --------------------------------------------------------
    def construir_view_orcamento():
        repopular_linhas_produto()

        cabecalho = ft.Column(
            controls=[
                ft.Text("Novo Orçamento", size=24, weight=ft.FontWeight.BOLD, color=TEXTO_PRIMARIO),
                ft.Text(
                    "Preencha os dados abaixo para gerar um orçamento",
                    size=13,
                    color=TEXTO_SECUNDARIO,
                ),
            ],
            spacing=2,
        )

        conteudo_central = ft.Column(
            controls=[
                ft.Container(content=cabecalho, padding=ft.Padding.only(left=36, right=36, top=30, bottom=10)),
                ft.Container(
                    content=ft.ResponsiveRow(
                        controls=[
                            ft.Container(content=card_formulario, col={"sm": 12, "md": 7}),
                            ft.Container(content=card_resultado, col={"sm": 12, "md": 5}),
                        ],
                    ),
                    padding=ft.Padding.symmetric(horizontal=36),
                ),
                ft.Container(content=barra_acoes, padding=ft.Padding.only(left=36, right=36, top=20, bottom=28)),
            ],
            scroll=ft.ScrollMode.AUTO,
            expand=True,
        )

        sidebar = construir_sidebar()

        return ft.View(
            route="/",
            padding=0,
            bgcolor=CINZA_FUNDO,
            controls=[
                ft.Row(
                    controls=[sidebar, ft.Container(content=conteudo_central, expand=True)],
                    expand=True,
                    spacing=0,
                )
            ],
        )

    def construir_sidebar():
        def item_nav(texto, rota, ativo=False):
            return ft.Container(
                content=ft.Text(
                    texto,
                    color=BRANCO,
                    weight=ft.FontWeight.BOLD if ativo else ft.FontWeight.NORMAL,
                    size=13,
                ),
                bgcolor=SIDEBAR_ITEM_ATIVO if ativo else "transparent",
                border_radius=8,
                padding=ft.Padding.symmetric(horizontal=14, vertical=10),
                margin=ft.Margin.symmetric(horizontal=14, vertical=4),
                on_click=lambda e: asyncio.create_task(page.push_route(rota)),
                ink=True,
            )

        rota_atual = page.route

        return ft.Container(
            width=240,
            bgcolor=SIDEBAR_BG,
            content=ft.Column(
                controls=[
                    ft.Container(
                        content=ft.Column(
                            controls=[
                                ft.Container(
                                    width=70,
                                    height=70,
                                    alignment=ft.Alignment.CENTER,
                                    content=ft.Image(
                                        src="assets/logo.png",
                                        width=60,
                                        height=60,
                                    )
                                ),
                                ft.Text(
                                    "LiftDesk",
                                    size=16,
                                    weight=ft.FontWeight.BOLD,
                                    color=BRANCO,
                                ),
                                ft.Text(
                                    "Sistema de Orçamentos",
                                    size=11,
                                    color="#E8C4C4",
                                ),
                            ],
                            horizontal_alignment=ft.CrossAxisAlignment.CENTER,
                            spacing=8,
                        ),
                        padding=ft.Padding.only(top=30, bottom=20),
                        alignment=ft.Alignment.CENTER,
                    ),
                    ft.Divider(color="#9B2C2C", height=1),
                    ft.Container(
                        content=ft.Text(
                            "MENU", size=11, weight=ft.FontWeight.BOLD, color="#E8C4C4"
                        ),
                        padding=ft.Padding.only(left=24, top=10, bottom=8),
                    ),
                    item_nav("Novo Orçamento", "/", ativo=(rota_atual == "/")),
                    item_nav("Produtos", "/produtos", ativo=(rota_atual == "/produtos")),
                    item_nav("Histórico", "/historico", ativo=(rota_atual == "/historico")),
                    item_nav("Contratos", "/contratos", ativo=(rota_atual == "/contratos")),
                    ft.Container(expand=True),
                    ft.Container(
                        content=ft.Text(
                            "v1.0 · uso interno", size=10, color="#C98A8A"
                        ),
                        padding=ft.Padding.only(left=20, bottom=24),
                    ),
                ],
                spacing=0,
                expand=True,
            ),
        )

    def limpar_historico():
        cursor.execute(
            "DELETE FROM orcamento_itens WHERE orcamento_id IN (SELECT id FROM orcamentos)"
        )
        cursor.execute("DELETE FROM orcamentos")
        conexao.commit()

    def construir_view_historico():
        cursor.execute(
            """
            SELECT id, data_hora, cliente, valor_total
            FROM orcamentos
            ORDER BY id DESC
            """
        )
        resultados = cursor.fetchall()

        async def confirmar_limpeza_historico(e):
            def confirmar(e):
                page.pop_dialog()
                limpar_historico()
                page.views.clear()
                page.views.append(construir_view_historico())
                page.update()

            def cancelar(e):
                page.pop_dialog()

            dialogo = ft.AlertDialog(
                modal=True,
                title=ft.Text("Limpar histórico"),
                content=ft.Text(
                    "Deseja realmente apagar todos os orçamentos salvos? "
                    "Essa ação não pode ser desfeita."
                ),
                actions=[
                    ft.TextButton("Cancelar", on_click=cancelar),
                    ft.TextButton(
                        "Limpar",
                        on_click=confirmar,
                        style=ft.ButtonStyle(color=VERMELHO),
                    ),
                ],
                actions_alignment=ft.MainAxisAlignment.END,
            )
            page.show_dialog(dialogo)

        tem_historico = bool(resultados)
        linhas = []
        if not tem_historico:
            linhas.append(
                ft.Container(
                    content=ft.Column(
                        controls=[
                            ft.Icon(
                                ft.Icons.HISTORY_EDU,
                                size=44,
                                color=VERMELHO,
                            ),
                            ft.Text(
                                "Nenhum orçamento registrado ainda.",
                                size=15,
                                weight=ft.FontWeight.BOLD,
                                color=TEXTO_PRIMARIO,
                            ),
                            ft.Text(
                                "Os orçamentos salvos aparecerão aqui com resumo e valor.",
                                size=12,
                                color=TEXTO_SECUNDARIO,
                                text_align=ft.TextAlign.CENTER,
                            ),
                        ],
                        horizontal_alignment=ft.CrossAxisAlignment.CENTER,
                        spacing=8,
                    ),
                    bgcolor=CINZA_CARD,
                    border_radius=16,
                    border=ft.Border.all(1, CINZA_BORDA),
                    padding=32,
                    alignment=ft.Alignment.CENTER,
                )
            )
        else:
            for item in resultados:
                id_, data_hora, cliente, valor_total = item

                cursor.execute(
                    "SELECT produto, quantidade FROM orcamento_itens WHERE orcamento_id = ?",
                    (id_,),
                )
                itens_do_orcamento = cursor.fetchall()
                if itens_do_orcamento:
                    resumo_produtos = ", ".join(
                        f"{qtd}x {nome}" for nome, qtd in itens_do_orcamento
                    )
                else:
                    resumo_produtos = "sem produtos"

                texto_principal = f"{cliente}  ·  {resumo_produtos}"
                texto_secundario = f"ID {id_}  ·  {data_hora}"
                try:
                    valor_formatado = formatar_moeda(valor_total)
                except (ValueError, TypeError):
                    # Registro antigo com dado inconsistente no banco —
                    # não deixa um item ruim quebrar a tela inteira.
                    valor_formatado = "valor inválido"

                linhas.append(
                    ft.Container(
                        content=ft.Row(
                            controls=[
                                ft.Container(
                                    content=ft.Icon(
                                        ft.Icons.RECEIPT_LONG,
                                        color=VERMELHO,
                                        size=20,
                                    ),
                                    bgcolor="#FDECEC",
                                    border_radius=999,
                                    padding=8,
                                    margin=ft.Margin.only(right=12),
                                ),
                                ft.Column(
                                    controls=[
                                        ft.Text(
                                            texto_principal,
                                            size=13,
                                            weight=ft.FontWeight.BOLD,
                                            color=TEXTO_PRIMARIO,
                                        ),
                                        ft.Text(
                                            texto_secundario,
                                            size=11,
                                            color=TEXTO_SECUNDARIO,
                                        ),
                                    ],
                                    spacing=2,
                                    expand=True,
                                ),
                                ft.Container(
                                    content=ft.Text(
                                        valor_formatado,
                                        size=13,
                                        weight=ft.FontWeight.BOLD,
                                        color=VERMELHO,
                                    ),
                                    bgcolor="#FFF5F5",
                                    border_radius=999,
                                    padding=ft.Padding.symmetric(horizontal=12, vertical=6),
                                ),
                            ],
                            alignment=ft.MainAxisAlignment.SPACE_BETWEEN,
                            vertical_alignment=ft.CrossAxisAlignment.START,
                        ),
                        bgcolor=CINZA_CARD,
                        border_radius=12,
                        border=ft.Border.all(1, CINZA_BORDA),
                        padding=14,
                        margin=ft.Margin.symmetric(vertical=5),
                    )
                )

        cabecalho_controles = [
            ft.Row(
                controls=[
                    ft.IconButton(
                        icon=ft.Icons.ARROW_BACK,
                        icon_color=BRANCO,
                        on_click=lambda e: asyncio.create_task(page.push_route("/")),
                    ),
                    ft.Text(
                        "HISTÓRICO DE ORÇAMENTOS",
                        size=18,
                        weight=ft.FontWeight.BOLD,
                        color=BRANCO,
                    ),
                ],
                vertical_alignment=ft.CrossAxisAlignment.CENTER,
            )
        ]
        if tem_historico:
            cabecalho_controles.append(
                ft.TextButton(
                    "Limpar tudo",
                    icon=ft.Icons.DELETE_SWEEP,
                    on_click=confirmar_limpeza_historico,
                    style=ft.ButtonStyle(
                        color=BRANCO,
                        bgcolor="#8B0000",
                        overlay_color="#A30000",
                    ),
                )
            )

        cabecalho_hist = ft.Container(
            content=ft.Row(
                controls=cabecalho_controles,
                alignment=ft.MainAxisAlignment.SPACE_BETWEEN,
                vertical_alignment=ft.CrossAxisAlignment.CENTER,
            ),
            bgcolor=VERMELHO,
            padding=ft.Padding.symmetric(horizontal=14, vertical=18),
        )

        lista = ft.Column(controls=linhas, scroll=ft.ScrollMode.AUTO, expand=True, spacing=8)

        conteudo_central = ft.Column(
            controls=[
                cabecalho_hist,
                ft.Container(content=lista, padding=20, expand=True),
            ],
            expand=True,
        )

        sidebar = construir_sidebar()

        return ft.View(
            route="/historico",
            padding=0,
            bgcolor=CINZA_FUNDO,
            controls=[
                ft.Row(
                    controls=[sidebar, ft.Container(content=conteudo_central, expand=True)],
                    expand=True,
                    spacing=0,
                )
            ],
        )

    def construir_view_produtos():
        produtos_atuais = carregar_produtos()

        async def confirmar_exclusao(nome):
            async def excluir_e_fechar(e):
                page.pop_dialog()
                page.update()
                await asyncio.sleep(0)
                excluir_produto(nome)
                ao_mudar_rota()  # reconstrói a view de produtos sem o item excluído

            def fechar_dialogo(e):
                page.pop_dialog()

            dialogo = ft.AlertDialog(
                modal=True,
                title=ft.Text("Confirmar exclusão"),
                content=ft.Text(
                    f"Tem certeza que deseja excluir o produto \"{nome}\"? "
                    f"Essa ação não pode ser desfeita."
                ),
                actions=[
                    ft.TextButton("Cancelar", on_click=fechar_dialogo),
                    ft.TextButton(
                        "Excluir",
                        on_click=excluir_e_fechar,
                        style=ft.ButtonStyle(color=VERMELHO),
                    ),
                ],
                actions_alignment=ft.MainAxisAlignment.END,
            )
            page.show_dialog(dialogo)

        linhas = []
        if not produtos_atuais:
            linhas.append(
                ft.Container(
                    content=ft.Text(
                        "Nenhum produto cadastrado ainda.",
                        size=13,
                        color=TEXTO_SECUNDARIO,
                    ),
                    padding=30,
                    alignment=ft.Alignment.CENTER,
                )
            )
        else:
            for nome, info in produtos_atuais.items():
                try:
                    preco_formatado = formatar_moeda(info["preco"])
                except (ValueError, TypeError):
                    preco_formatado = "valor inválido"

                caracteristicas = info.get("caracteristicas") or ""
                linha_caracteristicas = []
                if caracteristicas.strip():
                    linha_caracteristicas.append(
                        ft.Text(
                            caracteristicas,
                            size=11,
                            color=TEXTO_SECUNDARIO,
                        )
                    )

                linhas.append(
                    ft.Container(
                        content=ft.Row(
                            controls=[
                                ft.Column(
                                    controls=[
                                        ft.Text(
                                            nome,
                                            size=13,
                                            weight=ft.FontWeight.BOLD,
                                            color=TEXTO_PRIMARIO,
                                        ),
                                        *linha_caracteristicas,
                                    ],
                                    spacing=2,
                                    expand=True,
                                ),
                                ft.Text(
                                    preco_formatado,
                                    size=14,
                                    weight=ft.FontWeight.BOLD,
                                    color=VERMELHO,
                                ),
                                ft.IconButton(
                                    icon=ft.Icons.DELETE_OUTLINE,
                                    icon_color=VERMELHO,
                                    tooltip="Excluir produto",
                                    on_click=lambda e, n=nome: asyncio.create_task(confirmar_exclusao(n)),
                                ),
                            ],
                            alignment=ft.MainAxisAlignment.SPACE_BETWEEN,
                        ),
                        bgcolor=CINZA_CARD,
                        border_radius=10,
                        border=ft.Border.all(1, CINZA_BORDA),
                        padding=14,
                        margin=ft.Margin.symmetric(vertical=5),
                    )
                )

        cabecalho_prod = ft.Container(
            content=ft.Row(
                controls=[
                    ft.Text(
                        "PRODUTOS CADASTRADOS",
                        size=18,
                        weight=ft.FontWeight.BOLD,
                        color=BRANCO,
                        expand=True,
                    ),
                    ft.ElevatedButton(
                        content="+ Novo Produto",
                        on_click=lambda e: asyncio.create_task(
                            page.push_route("/produtos/novo")
                        ),
                        style=ft.ButtonStyle(
                            bgcolor=BRANCO,
                            color=VERMELHO,
                            shape=ft.RoundedRectangleBorder(radius=10),
                        ),
                    ),
                ],
            ),
            bgcolor=VERMELHO,
            padding=ft.Padding.symmetric(horizontal=20, vertical=18),
        )

        lista = ft.Column(controls=linhas, scroll=ft.ScrollMode.AUTO, expand=True)

        conteudo_central = ft.Column(
            controls=[
                cabecalho_prod,
                ft.Container(content=lista, padding=20, expand=True),
            ],
            expand=True,
        )

        sidebar = construir_sidebar()

        return ft.View(
            route="/produtos",
            padding=0,
            bgcolor=CINZA_FUNDO,
            controls=[
                ft.Row(
                    controls=[sidebar, ft.Container(content=conteudo_central, expand=True)],
                    expand=True,
                    spacing=0,
                )
            ],
        )
    
    def construir_view_novo_produto():
        nome_produto_field = ft.TextField(
            label="Nome do produto",
            hint_text="Ex.: elevador hospitalar",
            height=56,
            border_radius=8,
            border_color=CINZA_BORDA,
            bgcolor="#FAFAFA",
            color=TEXTO_PRIMARIO,
            label_style=ft.TextStyle(color=TEXTO_SECUNDARIO),
            hint_style=ft.TextStyle(color=TEXTO_SECUNDARIO),
        )

        preco_produto_field = ft.TextField(
            label="Preço (R$)",
            hint_text="Ex.: 150000",
            height=56,
            border_radius=8,
            border_color=CINZA_BORDA,
            bgcolor="#FAFAFA",
            color=TEXTO_PRIMARIO,
            label_style=ft.TextStyle(color=TEXTO_SECUNDARIO),
            hint_style=ft.TextStyle(color=TEXTO_SECUNDARIO),
            keyboard_type=ft.KeyboardType.NUMBER,
        )

        caracteristicas_produto_field = ft.TextField(
            label="Características técnicas",
            hint_text=(
                "Ex.: Capacidade para 08 passageiros — 630Kg. Velocidade: 1,5m/s. "
                "Dimensões internas da cabine: 1.100L x 1.400P x 2.200mm."
            ),
            multiline=True,
            min_lines=4,
            max_lines=8,
            border_radius=8,
            border_color=CINZA_BORDA,
            bgcolor="#FAFAFA",
            color=TEXTO_PRIMARIO,
            label_style=ft.TextStyle(color=TEXTO_SECUNDARIO),
            hint_style=ft.TextStyle(color=TEXTO_SECUNDARIO),
        )

        erro_produto_text = ft.Container(
            content=ft.Text("", color=VERMELHO, weight=ft.FontWeight.BOLD, size=12),
            bgcolor="#FCEAEA",
            border_radius=8,
            padding=8,
            visible=False,
        )

        def mostrar_erro_produto(mensagem):
            erro_produto_text.content.value = mensagem
            erro_produto_text.visible = True
            page.update()

        async def salvar_novo_produto(e):
            nome = (nome_produto_field.value or "").strip()
            preco_texto = (preco_produto_field.value or "").strip()

            if nome == "":
                mostrar_erro_produto("Digite o nome do produto")
                return

            if produto_existe(nome):
                mostrar_erro_produto("Já existe um produto com esse nome")
                return

            preco_texto_normalizado = preco_texto.replace(",", ".")
            try:
                preco = float(preco_texto_normalizado)
                if preco <= 0:
                    raise ValueError
            except ValueError:
                mostrar_erro_produto("Digite um preço válido (ex.: 150000)")
                return

            inserir_produto(nome, preco, caracteristicas_produto_field.value or "")
            mostrar_snack(f"Produto \"{nome}\" cadastrado com sucesso!")
            await page.push_route("/produtos")

        form_card = ft.Container(
            content=ft.Column(
                controls=[
                    nome_produto_field,
                    preco_produto_field,
                    caracteristicas_produto_field,
                    erro_produto_text,
                    ft.Row(
                        controls=[
                            ft.ElevatedButton(
                                content="Salvar Produto",
                                on_click=salvar_novo_produto,
                                style=ft.ButtonStyle(
                                    bgcolor=VERMELHO,
                                    color=BRANCO,
                                    shape=ft.RoundedRectangleBorder(radius=10),
                                ),
                            ),
                            ft.ElevatedButton(
                                content="Cancelar",
                                on_click=lambda e: asyncio.create_task(
                                    page.push_route("/produtos")
                                ),
                                style=ft.ButtonStyle(
                                    bgcolor=BRANCO,
                                    color=TEXTO_PRIMARIO,
                                    shape=ft.RoundedRectangleBorder(radius=10),
                                    side=ft.BorderSide(1, CINZA_BORDA),
                                ),
                            ),
                        ],
                        spacing=10,
                    ),
                ],
                spacing=16,
            ),
            bgcolor=CINZA_CARD,
            border_radius=14,
            border=ft.Border.all(1, CINZA_BORDA),
            padding=26,
            width=480,
        )

        cabecalho_novo = ft.Column(
            controls=[
                ft.Text("Novo Produto", size=24, weight=ft.FontWeight.BOLD, color=TEXTO_PRIMARIO),
                ft.Text(
                    "Cadastre um novo produto para usar nos orçamentos",
                    size=13,
                    color=TEXTO_SECUNDARIO,
                ),
            ],
            spacing=2,
        )

        conteudo_central = ft.Column(
            controls=[
                ft.Container(
                    content=cabecalho_novo,
                    padding=ft.Padding.only(left=36, right=36, top=30, bottom=10),
                ),
                ft.Container(
                    content=form_card,
                    padding=ft.Padding.symmetric(horizontal=36),
                ),
            ],
            scroll=ft.ScrollMode.AUTO,
            expand=True,
        )

        sidebar = construir_sidebar()

        return ft.View(
            route="/produtos/novo",
            padding=0,
            bgcolor=CINZA_FUNDO,
            controls=[
                ft.Row(
                    controls=[sidebar, ft.Container(content=conteudo_central, expand=True)],
                    expand=True,
                    spacing=0,
                )
            ],
        )

    def construir_view_contratos():
        orcamentos_salvos = carregar_lista_orcamentos()

        estado_contrato = {"orcamento_id": None, "dados": None}

        opcoes_orcamento = [
            ft.DropdownOption(
                key=str(id_),
                text=f"#{id_} — {cliente} — {data_hora} — {formatar_moeda_seguro(valor_total or 0)}",
            )
            for id_, data_hora, cliente, valor_total in orcamentos_salvos
        ]

        orcamento_dropdown = ft.Dropdown(
            label="Selecionar Orçamento",
            width=520,
            options=opcoes_orcamento,
            border_radius=8,
            bgcolor=VERMELHO,
            color=TEXTO_PRIMARIO,
            height=56,
        )

        cliente_contrato_field = ft.TextField(
            label="Cliente",
            read_only=True,
            width=520,
            height=56,
            border_radius=8,
            
            border_color=CINZA_BORDA,
            bgcolor="#FFFBFB",
            color=TEXTO_PRIMARIO,
            label_style=ft.TextStyle(color=TEXTO_SECUNDARIO),
        )
        nome_da_empresa_field = ft.TextField(
            label="Nome da Empresa Contratante",
            hint_text="Ex.: Empresa ABC Ltda.",
            width=520,
            expand=True,
            height=56,
            border_radius=8,
            border_color=CINZA_BORDA,
            bgcolor="#FAFAFA",
            color=TEXTO_PRIMARIO,
            label_style=ft.TextStyle(color=TEXTO_SECUNDARIO),
        )
        cnpj_field = ft.TextField(
            label="CNPJ/CPF do Cliente",
            hint_text="00.000.000/0000-00",
            width=400,
            height=56,
            border_radius=8,
            border_color=CINZA_BORDA,
            bgcolor="#FAFAFA",
            color=TEXTO_PRIMARIO,
            label_style=ft.TextStyle(color=TEXTO_SECUNDARIO),
            hint_style=ft.TextStyle(color=TEXTO_SECUNDARIO),
        )
        razao_social_field = ft.TextField(
            label="Razão Social da Empresa",
            width=520,
            expand=True,
            height=56,
            border_radius=8,
            border_color=CINZA_BORDA,
            bgcolor="#FAFAFA",
            color=TEXTO_PRIMARIO,
            label_style=ft.TextStyle(color=TEXTO_SECUNDARIO)
        )
        endereco_empresa_field = ft.TextField(
            label="Endereço da Empresa",
            width=520,
            expand=True,
            height=56,
            border_radius=8,
            border_color=CINZA_BORDA,
            bgcolor="#FAFAFA",
            color=TEXTO_PRIMARIO,
            label_style=ft.TextStyle(color=TEXTO_SECUNDARIO)
        )
        nome_do_representante_field = ft.TextField(
            label="Nome do Representante",
            width=520,
            expand=True,
            height=56,
            border_radius=8,
            border_color=CINZA_BORDA,
            bgcolor="#FAFAFA",
            color=TEXTO_PRIMARIO,
            label_style=ft.TextStyle(color=TEXTO_SECUNDARIO),
        )
        nacionalidade_representante_field = ft.TextField(
            label="Nacionalidade do Representante",
            width=520,
            expand=True,
            height=56,
            border_radius=8,
            border_color=CINZA_BORDA,
            bgcolor="#FAFAFA",
            color=TEXTO_PRIMARIO,
            label_style=ft.TextStyle(color=TEXTO_SECUNDARIO),
        )
        estado_civil_representante_field = ft.TextField(
            label="Estado Civil do Representante",
            width=520,
            expand=True,
            height=56,
            border_radius=8,
            border_color=CINZA_BORDA,
            bgcolor="#FAFAFA",
            color=TEXTO_PRIMARIO,
            label_style=ft.TextStyle(color=TEXTO_SECUNDARIO),
        )
        regime_casamento_representante_field = ft.TextField(
            label="Regime do Casamento",
            width=520,
            expand=True,
            height=56,
            border_radius=8,
            border_color=CINZA_BORDA,
            bgcolor="#FAFAFA",
            color=TEXTO_PRIMARIO,
            label_style=ft.TextStyle(color=TEXTO_SECUNDARIO),
        )
        numero_rg_representante_field = ft.TextField(
            label="Número do RG",
            width=520,
            expand=True,
            height=56,
            border_radius=8,
            border_color=CINZA_BORDA,
            bgcolor="#FAFAFA",
            color=TEXTO_PRIMARIO,
            label_style=ft.TextStyle(color=TEXTO_SECUNDARIO),
        )
        cpf_representante_field = ft.TextField(
            label="CPF do Representante",
            width=520,
            expand=True,
            height=56,
            border_radius=8,
            border_color=CINZA_BORDA,
            bgcolor="#FAFAFA",
            color=TEXTO_PRIMARIO,
            label_style=ft.TextStyle(color=TEXTO_SECUNDARIO),
        )
        endereco_representante_field = ft.TextField(
            label="Endereço do Representante",
            width=520,
            expand=True,
            height=56,
            border_radius=8,
            border_color=CINZA_BORDA,
            bgcolor="#FAFAFA",
            color=TEXTO_PRIMARIO,
            label_style=ft.TextStyle(color=TEXTO_SECUNDARIO),
        )

        resumo_contrato_text = ft.Text(
            "",
            size=13,
            color=TEXTO_SECUNDARIO,
            selectable=True,
        )

        erro_contrato = ft.Container(
            content=ft.Text("", color=VERMELHO, weight=ft.FontWeight.BOLD, size=12),
            bgcolor="#FCEAEA",
            border_radius=8,
            padding=8,
            visible=False,
        )

        def mostrar_erro_contrato(mensagem):
            erro_contrato.content.value = mensagem
            erro_contrato.visible = True
            page.update()

        def esconder_erro_contrato():
            erro_contrato.visible = False

        def atualizar_estado_do_orcamento():
            valor_selecionado = orcamento_dropdown.value
            if not valor_selecionado:
                estado_contrato["orcamento_id"] = None
                estado_contrato["dados"] = None
                cliente_contrato_field.value = ""
                resumo_contrato_text.value = ""
                return None

            try:
                orcamento_id = int(valor_selecionado)
            except (TypeError, ValueError):
                estado_contrato["orcamento_id"] = None
                estado_contrato["dados"] = None
                cliente_contrato_field.value = ""
                resumo_contrato_text.value = ""
                return None

            dados = carregar_orcamento_completo(orcamento_id)
            estado_contrato["orcamento_id"] = orcamento_id
            estado_contrato["dados"] = dados

            if dados:
                cliente_contrato_field.value = dados["cliente"]
                resumo_contrato_text.value = resumo_contrato(dados)
            else:
                cliente_contrato_field.value = ""
                resumo_contrato_text.value = ""
            return dados

        def ao_selecionar_orcamento(e):
            esconder_erro_contrato()
            atualizar_estado_do_orcamento()
            page.update()

        orcamento_dropdown.on_change = ao_selecionar_orcamento
        def resumo_contrato(dados):
            if not dados:
                dados = estado_contrato.get("dados")
            if dados:
                resumo_produtos = "\n".join(
                    f"{item['quantidade']}x {item['produto']} — "
                    f"{formatar_moeda(item['preco_unitario'])} cada"
                    for item in dados["itens"]
                )
                return (
                    f"Produtos:\n{resumo_produtos}\n\n"
                    f"Condição de Pagamento: {dados['parcela']}\n"
                    f"Valor Total: {dados['valor_total_formatado']}\n"
                    f"Valor da Parcela: {dados['valor_parcela_formatado']}"
                )
            return ""

        async def salvar_contrato_word(e):
            dados = estado_contrato.get("dados")
            if not dados:
                dados = atualizar_estado_do_orcamento()

            if not dados:
                mostrar_erro_contrato("Selecione um orçamento primeiro")
                return

            cnpj = (cnpj_field.value or "").strip()
            if not cnpj:
                mostrar_erro_contrato("Digite o CNPJ/CPF do cliente")
                return
            
            nome_empresa_contratante = (nome_da_empresa_field.value or "").strip()
            esconder_erro_contrato()

            placeholders = {
                "CONTRATANTE_NOME": nome_empresa_contratante or dados.get("cliente", ""),
                "CONTRATANTE_CNPJ": cnpj,
                "CONTRATANTE_ENDERECO": endereco_empresa_field.value or "",
                "CONTRATANTE_REPRESENTANTE": nome_do_representante_field.value or "",
                "CONTRATANTE_REPRESENTANTE_RG": numero_rg_representante_field.value or "",
                "CONTRATANTE_REPRESENTANTE_CPF": cpf_representante_field.value or "",
                "CONTRATANTE_REPRESENTANTE_ENDERECO": endereco_representante_field.value or "",
                "CONTRATANTE_REPRESENTANTE_NACIONALIDADE": nacionalidade_representante_field.value or "",
                "CONTRATANTE_REPRESENTANTE_ESTADO_CIVIL": estado_civil_representante_field.value or "",
                "CONTRATANTE_REPRESENTANTE_REGIME_CASAMENTO": regime_casamento_representante_field.value or "",
                "VALOR_TOTAL": dados.get("valor_total_formatado", ""),
                "VALOR_SINAL": dados.get("valor_parcela_formatado", ""),
                "VALOR_SINAL_EXTENSO": dados.get("valor_parcela_formatado", ""),
                "VALOR_PARCELA": dados.get("valor_parcela_formatado", ""),
                "PRAZO_ENTREGA": "90 (noventa) dias",
                "DATA_ASSINATURA": datetime.now().strftime("%d de %B de %Y"),
                "REPRESENTANTE_EMAIL": dados.get("email", ""),
            }

            nome_arquivo = f"contrato_{dados['cliente']}".strip().replace(" ", "_") + ".docx"
            caminho = await file_picker.save_file(
                dialog_title="Salvar contrato em Word",
                file_name=nome_arquivo,
                allowed_extensions=["docx"],
            )
            if not caminho:
                return
            if not caminho.lower().endswith(".docx"):
                caminho += ".docx"

            try:
                gerar_contrato_word(caminho, placeholders)
            except Exception as erro:
                mostrar_erro_contrato(f"Erro ao gerar contrato: {erro}")
                return

            registrar_contrato(dados["id"], dados["cliente"], cnpj)
            mostrar_snack(f"Contrato salvo em: {caminho}")

        card_contrato = ft.Container(
            content=ft.Column(
                controls=[
                    orcamento_dropdown,
                    cliente_contrato_field,
                    cnpj_field,
                    nome_da_empresa_field,
                    erro_contrato,
                    razao_social_field,
                    endereco_empresa_field,
                    nome_do_representante_field,
                    nacionalidade_representante_field,
                    estado_civil_representante_field,
                    regime_casamento_representante_field,
                    numero_rg_representante_field,
                    cpf_representante_field,
                    endereco_representante_field,
                    ft.Row(
                        controls=[
                            ft.ElevatedButton(
                                content="Salvar Contrato em Word",
                                on_click=salvar_contrato_word,
                                style=ft.ButtonStyle(
                                    bgcolor=VERMELHO,
                                    color=BRANCO,
                                    shape=ft.RoundedRectangleBorder(radius=10),
                                ),
                            ),
                        ],
                        spacing=10,
                    ),
                ],
                spacing=16,
            ),
            bgcolor=CINZA_CARD,
            border_radius=14,
            border=ft.Border.all(1, CINZA_BORDA),
            padding=26,
        )

        card_resumo_contrato = ft.Container(
            content=ft.Column(
                controls=[
                    ft.Text(
                        "RESUMO DO ORÇAMENTO SELECIONADO",
                        size=12,
                        weight=ft.FontWeight.BOLD,
                        color=TEXTO_SECUNDARIO,
                    ),
                    resumo_contrato_text,
                ],
                spacing=14,
            ),
            bgcolor=CINZA_CARD,
            border_radius=14,
            border=ft.Border.all(1, CINZA_BORDA),
            padding=24,
        )

        cabecalho_contrato = ft.Column(
            controls=[
                ft.Text("Novo Contrato", size=24, weight=ft.FontWeight.BOLD, color=TEXTO_PRIMARIO),
                ft.Text(
                    "Escolha um orçamento salvo para gerar o contrato em Word",
                    size=13,
                    color=TEXTO_SECUNDARIO,
                ),
            ],
            spacing=2,
        )

        aviso_sem_orcamentos = []
        if not orcamentos_salvos:
            aviso_sem_orcamentos.append(
                ft.Container(
                    content=ft.Text(
                        "Nenhum orçamento salvo ainda. Crie um orçamento primeiro "
                        "na tela \"Novo Orçamento\".",
                        size=13,
                        color=TEXTO_SECUNDARIO,
                    ),
                    padding=ft.Padding.symmetric(horizontal=36),
                )
            )

        conteudo_central = ft.Column(
            controls=[
                ft.Container(
                    content=cabecalho_contrato,
                    padding=ft.Padding.only(left=36, right=36, top=30, bottom=10),
                ),
                *aviso_sem_orcamentos,
                ft.Container(
                    content=ft.ResponsiveRow(
                        controls=[
                            ft.Container(content=card_contrato, col={"sm": 12, "md": 7}),
                            ft.Container(content=card_resumo_contrato, col={"sm": 12, "md": 5}),
                        ],
                    ),
                    padding=ft.Padding.symmetric(horizontal=36),
                ),
            ],
            scroll=ft.ScrollMode.AUTO,
            expand=True,
        )

        sidebar = construir_sidebar()

        return ft.View(
            route="/contratos",
            padding=0,
            bgcolor=CINZA_FUNDO,
            controls=[
                ft.Row(
                    controls=[sidebar, ft.Container(content=conteudo_central, expand=True)],
                    expand=True,
                    spacing=0,
                )
            ],
        )

    def ao_mudar_rota(e=None):
        page.views.clear()
        if page.route == "/historico":
            page.views.append(construir_view_historico())
        elif page.route == "/produtos/novo":
            page.views.append(construir_view_novo_produto())
        elif page.route == "/produtos":
            page.views.append(construir_view_produtos())
        elif page.route == "/contratos":
            page.views.append(construir_view_contratos())
        else:
            page.views.append(construir_view_orcamento())
        page.update()

    async def ao_voltar(e):
        if getattr(e, "view", None) is not None:
            page.views.remove(e.view)
        if page.views:
            top_view = page.views[-1]
            await page.push_route(top_view.route)

    page.on_route_change = ao_mudar_rota
    page.on_view_pop = ao_voltar

    ao_mudar_rota()

    def ao_fechar(e):
        conexao.close()

    page.on_disconnect = ao_fechar
    #banco de dados
    if produto_existe("Elevador Social") is False:
        inserir_produto(
            "Elevador Social",
            180000,
            "Capacidade para 08 passageiros - 630Kg. Velocidade: 1,5m/s. Dimensões internas da cabine: 1.100L x 1.400P x 2.200mm."
        )


if __name__ == "__main__":
    ft.run(main)
